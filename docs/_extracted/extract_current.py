"""Extract structure + text from the current First Term docx."""
import json
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

src = Path(r"D:/Courses/Level_4/Graduation Project/Code_Review_Platform/Code Mentor V1/docs/Decumentation finaly version-2.1.docx")
out_dir = Path(r"D:/Courses/Level_4/Graduation Project/Code_Review_Platform/Code Mentor V1/docs/_extracted")
out_dir.mkdir(parents=True, exist_ok=True)

doc = Document(str(src))

# Walk paragraphs in document order with their style and approximate level.
lines = []
toc = []
for i, p in enumerate(doc.paragraphs):
    style = (p.style.name if p.style else "Normal") or "Normal"
    text = p.text.strip()
    if not text:
        lines.append({"i": i, "style": style, "text": ""})
        continue
    lines.append({"i": i, "style": style, "text": text})
    if style.startswith("Heading"):
        toc.append({"i": i, "style": style, "text": text})

# Also walk tables — they may carry structure (WBS, Gantt) we care about
tables = []
for ti, t in enumerate(doc.tables):
    rows = []
    for r in t.rows:
        rows.append([c.text.strip() for c in r.cells])
    tables.append({"ti": ti, "rows": rows})

# Inline shapes (images) count + relationships
inline_shape_count = len(doc.inline_shapes)

# Save structured outputs
(out_dir / "current_toc.json").write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
(out_dir / "current_full.json").write_text(json.dumps(lines, ensure_ascii=False, indent=2), encoding="utf-8")
(out_dir / "current_tables.json").write_text(json.dumps(tables, ensure_ascii=False, indent=2), encoding="utf-8")
(out_dir / "current_full.txt").write_text(
    "\n".join(f"[{l['style']}] {l['text']}" for l in lines), encoding="utf-8"
)

print(f"Total paragraphs: {len(lines)}")
print(f"Headings: {len(toc)}")
print(f"Tables: {len(tables)}")
print(f"Inline shapes (images): {inline_shape_count}")
print()
print("=== TOC ===")
for h in toc:
    indent = "  " * (int(h["style"].split()[-1]) - 1 if h["style"][-1].isdigit() else 0)
    print(f"{indent}{h['text']}")
