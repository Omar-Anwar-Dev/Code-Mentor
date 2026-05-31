"""Extract structure of reference docs for format alignment."""
import json
from pathlib import Path
from docx import Document
import pdfplumber

out_dir = Path(r"D:/Courses/Level_4/Graduation Project/Code_Review_Platform/Code Mentor V1/docs/_extracted")

# --- PlantCare docx ---
plantcare = Path(r"D:/Courses/Level_4/Graduation Project/Code_Review_Platform/Code Mentor V1/docs/PlantCare_Documentation.docx")
doc = Document(str(plantcare))
pc_headings = []
for p in doc.paragraphs:
    style = (p.style.name if p.style else "Normal") or "Normal"
    text = p.text.strip()
    if style.startswith("Heading") or style.lower().startswith("toc"):
        if text:
            pc_headings.append({"style": style, "text": text})

(out_dir / "plantcare_headings.json").write_text(
    json.dumps(pc_headings, ensure_ascii=False, indent=2), encoding="utf-8"
)

# Tables in PlantCare
pc_tables = [{"ti": ti, "rows": [[c.text.strip() for c in r.cells] for r in t.rows]}
             for ti, t in enumerate(doc.tables)]
(out_dir / "plantcare_tables.json").write_text(
    json.dumps(pc_tables, ensure_ascii=False, indent=2), encoding="utf-8"
)

print(f"PlantCare paragraphs: {len(doc.paragraphs)}")
print(f"PlantCare headings: {len(pc_headings)}")
print(f"PlantCare tables: {len(pc_tables)}")
print(f"PlantCare images: {len(doc.inline_shapes)}")
print()

# --- First Term PDF ---
pdf_path = Path(r"D:/Courses/Level_4/Graduation Project/Code_Review_Platform/Code Mentor V1/docs/Graduation Documentation-First Term-Final File.pdf")
with pdfplumber.open(str(pdf_path)) as pdf:
    pdf_pages = len(pdf.pages)
    # Pull text only from TOC pages (usually 1-5)
    toc_text = ""
    for p in pdf.pages[:6]:
        text = p.extract_text() or ""
        toc_text += text + "\n---\n"
    (out_dir / "ref_pdf_toc_pages.txt").write_text(toc_text, encoding="utf-8")
    print(f"Ref PDF pages: {pdf_pages}")
print("Wrote: plantcare_headings.json, plantcare_tables.json, ref_pdf_toc_pages.txt")
